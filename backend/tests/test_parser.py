"""
文档解析器单元测试
"""
import os
import sys
import tempfile
import pytest
from pathlib import Path
from io import BytesIO

# 添加项目根目录到路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.processors.parser import DocumentParser, parse_document, parse_document_from_bytes


class TestDocumentParser:
    """文档解析器测试类"""

    @pytest.fixture
    def temp_dir(self):
        """创建临时目录"""
        with tempfile.TemporaryDirectory() as tmpdir:
            yield tmpdir

    # ========== 基础测试 ==========

    def test_supported_types(self):
        """测试支持的文档类型"""
        expected_types = [
            "pdf", "docx", "xlsx", "md", "txt",
            "pptx", "ppt", "doc", "html", "htm",
            "rtf", "png", "jpg", "jpeg", "gif",
            "bmp", "tiff", "webp"
        ]
        for file_type in expected_types:
            assert file_type in DocumentParser.SUPPORTED_TYPES

    def test_unsupported_type_raises_error(self):
        """测试不支持的文件类型抛出异常"""
        with pytest.raises(ValueError) as exc_info:
            DocumentParser.parse("/fake/path.xyz", "xyz")
        assert "Unsupported file type" in str(exc_info.value)

    # ========== 文本文件测试 ==========

    def test_parse_text_utf8(self, temp_dir):
        """测试解析UTF-8文本文件"""
        content = "Hello World\n这是测试内容"
        file_path = os.path.join(temp_dir, "test.txt")
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)

        result = DocumentParser.parse(file_path, "txt")

        assert result["content"] == content
        assert result["metadata"]["format"] == "TXT"
        assert result["metadata"]["page_count"] == 1
        assert len(result["pages"]) == 1
        assert result["pages"][0]["content"] == content

    def test_parse_text_gbk(self, temp_dir):
        """测试解析GBK编码文本文件"""
        content = "GBK编码测试"
        file_path = os.path.join(temp_dir, "test_gbk.txt")
        with open(file_path, "w", encoding="gbk") as f:
            f.write(content)

        result = DocumentParser.parse(file_path, "txt")

        assert content in result["content"]
        assert result["metadata"]["encoding"] in ["gbk", "utf-8"]

    def test_parse_text_empty(self, temp_dir):
        """测试解析空文本文件"""
        file_path = os.path.join(temp_dir, "empty.txt")
        Path(file_path).touch()

        result = DocumentParser.parse(file_path, "txt")

        assert result["content"] == ""
        assert result["metadata"]["page_count"] == 1

    # ========== Markdown测试 ==========

    def test_parse_markdown(self, temp_dir):
        """测试解析Markdown文件"""
        content = """# 标题

这是正文内容。

## 二级标题

- 列表项1
- 列表项2
"""
        file_path = os.path.join(temp_dir, "test.md")
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)

        result = DocumentParser.parse(file_path, "md")

        assert result["content"] == content
        assert result["metadata"]["format"] == "Markdown"
        assert result["metadata"]["page_count"] == 1

    # ========== RTF测试 ==========

    def test_parse_rtf(self, temp_dir):
        """测试解析RTF文件"""
        rtf_content = r"{\rtf1\ansi{\fonttbl\f0 Courier;}\f0\pard This is some {\b bold} text.\par}"
        file_path = os.path.join(temp_dir, "test.rtf")
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(rtf_content)

        result = DocumentParser.parse(file_path, "rtf")

        assert "bold" in result["content"]
        assert "text" in result["content"]
        assert result["metadata"]["format"] == "RTF"

    # ========== HTML测试 ==========

    def test_parse_html(self, temp_dir):
        """测试解析HTML文件"""
        html_content = """<!DOCTYPE html>
<html>
<head><title>Test Page</title></head>
<body>
<h1>Hello World</h1>
<p>This is a test paragraph.</p>
<a href="http://example.com">Link</a>
</body>
</html>
"""
        file_path = os.path.join(temp_dir, "test.html")
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        result = DocumentParser.parse(file_path, "html")

        assert "Hello World" in result["content"]
        assert "Test Page" in result["content"]
        assert result["metadata"]["format"] == "HTML"
        assert result["metadata"]["title"] == "Test Page"

    def test_parse_html_without_title(self, temp_dir):
        """测试解析无标题HTML"""
        html_content = "<html><body><p>No title here</p></body></html>"
        file_path = os.path.join(temp_dir, "test2.html")
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        result = DocumentParser.parse(file_path, "html")

        assert "No title here" in result["content"]
        assert result["metadata"]["title"] is None

    # ========== Word docx测试 ==========

    def test_parse_docx_basic(self, temp_dir):
        """测试解析基本docx文件"""
        try:
            from docx import Document

            # 创建一个简单的docx文件
            file_path = os.path.join(temp_dir, "test.docx")
            doc = Document()
            doc.add_heading("测试文档", 0)
            doc.add_paragraph("这是第一段内容。")
            doc.add_paragraph("这是第二段内容。")
            doc.save(file_path)

            result = DocumentParser.parse(file_path, "docx")

            assert "测试文档" in result["content"]
            assert "第一段" in result["content"]
            assert result["metadata"]["format"] == "DOCX"
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_parse_docx_with_table(self, temp_dir):
        """测试解析带表格的docx文件"""
        try:
            from docx import Document

            file_path = os.path.join(temp_dir, "test_table.docx")
            doc = Document()
            doc.add_paragraph("文档正文")

            # 添加表格
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "A1"
            table.cell(0, 1).text = "B1"
            table.cell(1, 0).text = "A2"
            table.cell(1, 1).text = "B2"
            doc.save(file_path)

            result = DocumentParser.parse(file_path, "docx")

            assert "文档正文" in result["content"]
            assert "A1" in result["content"]
            assert result["metadata"]["table_count"] == 1
        except ImportError:
            pytest.skip("python-docx not installed")

    # ========== Excel测试 ==========

    def test_parse_xlsx_basic(self, temp_dir):
        """测试解析基本xlsx文件"""
        try:
            from openpyxl import Workbook

            file_path = os.path.join(temp_dir, "test.xlsx")
            wb = Workbook()
            ws = wb.active
            ws.title = "Sheet1"
            ws["A1"] = "Name"
            ws["B1"] = "Age"
            ws["A2"] = "Alice"
            ws["B2"] = 25
            wb.save(file_path)

            result = DocumentParser.parse(file_path, "xlsx")

            assert "Name" in result["content"]
            assert "Alice" in result["content"]
            assert result["metadata"]["format"] == "XLSX"
            assert result["metadata"]["sheet_count"] == 1
        except ImportError:
            pytest.skip("openpyxl not installed")

    def test_parse_xlsx_multiple_sheets(self, temp_dir):
        """测试解析多工作表xlsx文件"""
        try:
            from openpyxl import Workbook

            file_path = os.path.join(temp_dir, "test_multi.xlsx")
            wb = Workbook()

            ws1 = wb.active
            ws1.title = "Sheet1"
            ws1["A1"] = "Data1"

            ws2 = wb.create_sheet("Sheet2")
            ws2["A1"] = "Data2"
            wb.save(file_path)

            result = DocumentParser.parse(file_path, "xlsx")

            assert "Data1" in result["content"]
            assert "Data2" in result["content"]
            assert result["metadata"]["sheet_count"] == 2
            assert "Sheet1" in result["metadata"]["sheet_names"]
            assert "Sheet2" in result["metadata"]["sheet_names"]
        except ImportError:
            pytest.skip("openpyxl not installed")

    # ========== PDF测试 ==========

    def test_parse_pdf_requires_file(self):
        """测试解析不存在的PDF文件抛出异常"""
        with pytest.raises(Exception):
            DocumentParser.parse("/nonexistent/file.pdf", "pdf")

    # ========== PPT测试 ==========

    def test_parse_pptx_requires_library(self):
        """测试PPT解析需要python-pptx库"""
        # 临时禁用pptx
        import app.processors.parser as parser_module
        original = parser_module.PPTX_AVAILABLE

        try:
            parser_module.PPTX_AVAILABLE = False

            with tempfile.NamedTemporaryFile(suffix=".pptx", delete=False) as tmp:
                tmp.write(b"fake pptx content")
                tmp_path = tmp.name

            try:
                with pytest.raises(ImportError):
                    DocumentParser.parse(tmp_path, "pptx")
            finally:
                os.unlink(tmp_path)
        finally:
            parser_module.PPTX_AVAILABLE = original

    # ========== 图片OCR测试 ==========

    def test_parse_image_requires_library(self):
        """测试图片解析需要pytesseract库"""
        import app.processors.parser as parser_module
        original = parser_module.PYTESSERACT_AVAILABLE

        try:
            parser_module.PYTESSERACT_AVAILABLE = False

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp.write(b"fake image content")
                tmp_path = tmp.name

            try:
                with pytest.raises(ImportError):
                    DocumentParser.parse(tmp_path, "png")
            finally:
                os.unlink(tmp_path)
        finally:
            parser_module.PYTESSERACT_AVAILABLE = original

    # ========== HTML解析测试 ==========

    def test_parse_html_requires_library(self):
        """测试HTML解析需要beautifulsoup库"""
        import app.processors.parser as parser_module
        original = parser_module.BEAUTIFULSOUP_AVAILABLE

        try:
            parser_module.BEAUTIFULSOUP_AVAILABLE = False

            with tempfile.NamedTemporaryFile(suffix=".html", delete=False) as tmp:
                tmp.write(b"<html><body>Test</body></html>")
                tmp_path = tmp.name

            try:
                with pytest.raises(ImportError):
                    DocumentParser.parse(tmp_path, "html")
            finally:
                os.unlink(tmp_path)
        finally:
            parser_module.BEAUTIFULSOUP_AVAILABLE = original

    # ========== 字节解析测试 ==========

    def test_parse_bytes(self, temp_dir):
        """测试从字节解析文档"""
        content = "测试从字节解析".encode("utf-8")
        file_type = "txt"

        result = parse_document_from_bytes(content, file_type)

        assert "测试从字节解析" in result["content"]
        assert result["metadata"]["format"] == "TXT"

    def test_parse_bytes_with_extension(self, temp_dir):
        """测试从字节解析带扩展名的文档"""
        content = b"Hello World"
        file_type = "txt"

        result = parse_document_from_bytes(content, file_type)

        assert result["content"] == "Hello World"

    # ========== 便捷函数测试 ==========

    def test_parse_document_function(self, temp_dir):
        """测试便捷函数"""
        content = "便捷函数测试"
        file_path = os.path.join(temp_dir, "test.txt")
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)

        result = parse_document(file_path, "txt")

        assert result["content"] == content

    # ========== 错误处理测试 ==========

    def test_parse_invalid_file_type(self):
        """测试无效文件类型"""
        with pytest.raises(ValueError) as exc_info:
            DocumentParser.parse("/fake/path.txt", "invalid_type")
        assert "Unsupported file type" in str(exc_info.value)

    def test_case_insensitive(self, temp_dir):
        """测试文件类型大小写不敏感"""
        content = "Case Test"
        file_path = os.path.join(temp_dir, "test.txt")
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)

        # 测试大写
        result = DocumentParser.parse(file_path, "TXT")
        assert result["content"] == content

        # 测试混合大小写
        result = DocumentParser.parse(file_path, "Txt")
        assert result["content"] == content


class TestDocumentParserEdgeCases:
    """边界情况测试"""

    @pytest.fixture
    def temp_dir(self):
        """创建临时目录"""
        with tempfile.TemporaryDirectory() as tmpdir:
            yield tmpdir

    def test_parse_file_with_spaces_in_name(self, temp_dir):
        """测试文件名包含空格"""
        content = "文件 名 含 空 格"
        file_path = os.path.join(temp_dir, "test file with spaces.txt")
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)

        result = DocumentParser.parse(file_path, "txt")

        assert result["content"] == content

    def test_parse_file_with_chinese_path(self, temp_dir):
        """测试中文路径"""
        content = "中文路径测试"
        # 创建带中文的目录和文件名
        chinese_dir = os.path.join(temp_dir, "测试目录")
        os.makedirs(chinese_dir, exist_ok=True)
        file_path = os.path.join(chinese_dir, "测试文件.txt")

        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)

        result = DocumentParser.parse(file_path, "txt")

        assert result["content"] == content


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
